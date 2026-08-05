#!/usr/bin/env python3
"""
One-off: builds a single combined .docx from
  1. docs/session-scheduling-use-cases.md (converted)
  2. docs/PWC-Backend-DB-Design-and-API-List.docx (appended as-is)
so it can be uploaded to Google Drive as one Google Doc.
Not part of the app runtime.
"""
import copy
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent.parent
MD_SRC = ROOT / "docs" / "session-scheduling-use-cases.md"
DOCX_SRC = ROOT / "docs" / "PWC-Backend-DB-Design-and-API-List.docx"
OUT = ROOT / "docs" / "PWC-Session-Scheduling-and-DB-Design-API-List.docx"


def add_inline_runs(paragraph, text):
    """Handle **bold** and `code` inline markers."""
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)


def convert_markdown(doc: Document, md_text: str):
    lines = md_text.split("\n")
    i = 0
    n = len(lines)
    in_code_block = False
    code_lines = []
    table_buffer = []

    def flush_table():
        if not table_buffer:
            return
        rows = [
            [c.strip() for c in row.strip().strip("|").split("|")]
            for row in table_buffer
            if not re.match(r"^\s*\|?\s*-{2,}", row)
        ]
        if not rows:
            table_buffer.clear()
            return
        ncols = len(rows[0])
        table = doc.add_table(rows=len(rows), cols=ncols)
        table.style = "Light Grid Accent 1"
        for r, row in enumerate(rows):
            for c, cell_text in enumerate(row):
                if c < ncols:
                    cell = table.cell(r, c)
                    cell.text = cell_text
                    if r == 0:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.bold = True
        table_buffer.clear()

    while i < n:
        line = lines[i]

        if line.strip().startswith("```"):
            if in_code_block:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                p.paragraph_format.space_after = Pt(8)
                code_lines = []
                in_code_block = False
            else:
                flush_table()
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        if line.strip().startswith("|"):
            table_buffer.append(line)
            i += 1
            continue
        else:
            flush_table()

        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._p.get_or_add_pPr()
            pBdr = pPr.makeelement(qn("w:pBdr"), {})
            bottom = pPr.makeelement(qn("w:bottom"), {
                qn("w:val"): "single", qn("w:sz"): "6", qn("w:space"): "1", qn("w:color"): "999999"
            })
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip("`")
            doc.add_heading(text, level=min(level, 4))
            i += 1
            continue

        m = re.match(r"^[-*]\s+(.*)$", stripped)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, m.group(1))
            i += 1
            continue

        m = re.match(r"^\d+\.\s+(.*)$", stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, m.group(1))
            i += 1
            continue

        p = doc.add_paragraph()
        add_inline_runs(p, stripped)
        i += 1

    flush_table()


def append_docx(dst: Document, src_path: Path):
    src = Document(str(src_path))
    body_children = list(src.element.body)
    for idx, element in enumerate(body_children):
        # Skip the trailing sectPr (section properties) — it belongs to src's own
        # section, not a paragraph; copying it in would create a spurious section
        # break. The final section of dst will carry its own sectPr already.
        if element.tag == qn("w:sectPr"):
            continue
        dst.element.body.append(copy.deepcopy(element))


def main():
    doc = Document()
    doc.add_heading(
        "Phoenix Water Club Counselling Platform — Session Scheduling & DB/API Reference",
        level=0,
    )
    doc.add_paragraph(
        "Combined export for sharing. Source files (kept up to date in the repo): "
        "docs/session-scheduling-use-cases.md and "
        "docs/PWC-Backend-DB-Design-and-API-List.docx."
    )

    doc.add_heading("Part 1: Session Scheduling — Use Cases & Flow", level=1)
    md_text = MD_SRC.read_text()
    # Drop the source file's own H1 (we already added our own top-level heading for this part).
    md_text = re.sub(r"^#\s+.*\n", "", md_text, count=1)
    convert_markdown(doc, md_text)

    # Page break before Part 2
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)

    doc.add_heading("Part 2: Database Design & API Reference", level=1)
    append_docx(doc, DOCX_SRC)

    doc.save(str(OUT))
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
